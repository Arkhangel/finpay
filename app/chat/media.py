"""MIME-диспатч медиавложений в OpenAI content-part."""

from __future__ import annotations

import base64
import logging
from io import BytesIO

from fastapi import UploadFile
from openai import AsyncOpenAI
from pypdf import PdfReader

from app.settings import settings

logger = logging.getLogger(__name__)

_MAX_PDF_PAGES = 50
_TEXT_TRUNCATE = 30_000
_SCAN_HEURISTIC_MIN_PAGES = 5
_SCAN_HEURISTIC_CHARS_PER_PAGE = 100


async def media_to_part(media: UploadFile, client: AsyncOpenAI) -> dict:
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"},
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        # gpt-*-audio's input_audio принимает только wav/mp3, а Telegram шлёт ogg/opus.
        transcript = await whisper_transcribe(client, data, media.filename or "audio.ogg")
        return {
            "type": "text",
            "text": f"[пользователь сказал голосом]:\n{transcript}",
        }

    if mime == "application/pdf":
        return {
            "type": "text",
            "text": f"[документ PDF]:\n{extract_pdf_text(data)[:_TEXT_TRUNCATE]}",
        }

    if mime.endswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{extract_docx_text(data)[:_TEXT_TRUNCATE]}",
        }

    raise ValueError(f"Unsupported media type: {mime}")


async def whisper_transcribe(client: AsyncOpenAI, audio_bytes: bytes, filename: str) -> str:
    """Whisper-1 принимает ogg/m4a/mp3/wav/flac/webm напрямую."""
    f = BytesIO(audio_bytes)
    f.name = filename  # SDK берёт расширение для определения формата
    result = await client.audio.transcriptions.create(
        model=settings.openai.whisper_model,
        file=f,
    )
    return result.text


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = reader.pages[:_MAX_PDF_PAGES]

    texts: list[str] = []
    short_pages = 0
    for page in pages:
        text = page.extract_text() or ""
        texts.append(text)
        if len(text.strip()) < _SCAN_HEURISTIC_CHARS_PER_PAGE:
            short_pages += 1

    if len(pages) >= _SCAN_HEURISTIC_MIN_PAGES and short_pages >= _SCAN_HEURISTIC_MIN_PAGES:
        logger.warning("pdf_looks_scanned pages=%d short_pages=%d", len(pages), short_pages)

    return "\n".join(texts)


def extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    chunks: list[str] = [p.text for p in doc.paragraphs if p.text]

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                chunks.append(row_text)

    return "\n".join(chunks)