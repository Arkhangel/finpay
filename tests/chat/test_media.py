"""Tests for MIME-dispatch in app/chat/media.py (no network calls for image/PDF)."""

from __future__ import annotations

import base64
from io import BytesIO

import pytest
from fastapi import UploadFile
from starlette.datastructures import Headers

from app.chat.media import extract_docx_text, extract_pdf_text, media_to_part

_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _upload_file(data: bytes, content_type: str, filename: str = "file.bin") -> UploadFile:
    return UploadFile(
        file=BytesIO(data),
        filename=filename,
        headers=Headers({"content-type": content_type}),
    )


def _blank_pdf_bytes() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


async def test_media_to_part_image_returns_data_uri():
    upload = _upload_file(_PNG_1X1, "image/png", "pic.png")

    part = await media_to_part(upload, client=None)

    assert part["type"] == "image_url"
    assert part["image_url"]["url"].startswith("data:image/png;base64,")
    b64_payload = part["image_url"]["url"].split(",", 1)[1]
    assert base64.b64decode(b64_payload) == _PNG_1X1


async def test_media_to_part_pdf_returns_text_part():
    upload = _upload_file(_blank_pdf_bytes(), "application/pdf", "doc.pdf")

    part = await media_to_part(upload, client=None)

    assert part["type"] == "text"
    assert part["text"].startswith("[документ PDF]:\n")


async def test_media_to_part_unsupported_mime_raises():
    upload = _upload_file(b"whatever", "application/zip", "archive.zip")

    with pytest.raises(ValueError, match="Unsupported media type"):
        await media_to_part(upload, client=None)


def test_extract_pdf_text_does_not_crash_on_blank_page():
    assert extract_pdf_text(_blank_pdf_bytes()) == ""


def test_extract_docx_text_reads_paragraphs_and_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Привет из DOCX")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"

    buf = BytesIO()
    doc.save(buf)

    text = extract_docx_text(buf.getvalue())

    assert "Привет из DOCX" in text
    assert "a\tb" in text
