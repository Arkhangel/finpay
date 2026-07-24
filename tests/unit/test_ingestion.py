"""Unit-тесты для app/services/ingestion.py: metadata-обогащение (блок 5.5).

Проверяются: разбор category из пути, version из имени файла, author из
DOCX core properties, и исключение шумных полей из embedding-текста.
"""
from __future__ import annotations

from docx import Document as DocxDocument
from llama_index.core.schema import Document

from app.services import ingestion


def test_file_metadata_extracts_category_from_subfolder(tmp_path):
    category_dir = tmp_path / "tariffs"
    category_dir.mkdir()
    file_path = category_dir / "01_fees.md"
    file_path.write_text("текст")

    file_metadata = ingestion.build_file_metadata(tmp_path)
    meta = file_metadata(str(file_path))

    assert meta["category"] == "tariffs"
    assert meta["source"] == "01_fees.md"


def test_file_metadata_category_none_for_flat_file(tmp_path):
    file_path = tmp_path / "01_fees.md"
    file_path.write_text("текст")

    file_metadata = ingestion.build_file_metadata(tmp_path)
    meta = file_metadata(str(file_path))

    assert meta["category"] is None


def test_file_metadata_extracts_version_from_filename(tmp_path):
    file_path = tmp_path / "legal" / "oferta_v2.docx"
    file_path.parent.mkdir()
    DocxDocument().save(file_path)

    file_metadata = ingestion.build_file_metadata(tmp_path)
    meta = file_metadata(str(file_path))

    assert meta["version"] == "2"


def test_file_metadata_version_none_without_suffix(tmp_path):
    file_path = tmp_path / "01_fees.md"
    file_path.write_text("текст")

    meta = ingestion.build_file_metadata(tmp_path)(str(file_path))

    assert meta["version"] is None


def test_file_metadata_extracts_docx_author(tmp_path):
    file_path = tmp_path / "legal" / "oferta.docx"
    file_path.parent.mkdir()
    doc = DocxDocument()
    doc.core_properties.author = "Иван Петров"
    doc.save(file_path)

    meta = ingestion.build_file_metadata(tmp_path)(str(file_path))

    assert meta["author"] == "Иван Петров"


def test_file_metadata_author_none_for_non_docx(tmp_path):
    file_path = tmp_path / "01_fees.md"
    file_path.write_text("текст")

    meta = ingestion.build_file_metadata(tmp_path)(str(file_path))

    assert meta["author"] is None


def test_file_metadata_includes_last_modified(tmp_path):
    file_path = tmp_path / "01_fees.md"
    file_path.write_text("текст")

    meta = ingestion.build_file_metadata(tmp_path)(str(file_path))

    assert "last_modified" in meta and meta["last_modified"]


def test_apply_embedding_exclusions_excludes_noisy_keys():
    doc = Document(
        text="текст документа",
        metadata={
            "source": "01_fees.md",
            "category": "tariffs",
            "last_modified": "2026-01-01T00:00:00",
            "version": "2",
        },
    )

    ingestion.apply_embedding_exclusions([doc])

    assert "last_modified" in doc.excluded_embed_metadata_keys
    assert "version" in doc.excluded_embed_metadata_keys
    assert "source" not in doc.excluded_embed_metadata_keys
    assert "category" not in doc.excluded_embed_metadata_keys


def test_apply_embedding_exclusions_skips_absent_keys():
    doc = Document(text="текст", metadata={"source": "01_fees.md"})

    ingestion.apply_embedding_exclusions([doc])

    assert doc.excluded_embed_metadata_keys == []
